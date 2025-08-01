from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import requests
from io import BytesIO
from pogoapi.utils.mongodb_client import db
from bson import ObjectId

def set_heading_style(doc):
    """Set up custom heading styles for SEO-friendly document structure."""
    # Define heading styles
    heading_styles = {
        'Heading 1': {'size': 24, 'color': RGBColor(0, 0, 0), 'bold': True},
        'Heading 2': {'size': 20, 'color': RGBColor(0, 0, 0), 'bold': True},
        'Heading 3': {'size': 16, 'color': RGBColor(0, 0, 0), 'bold': True},
        'Heading 4': {'size': 14, 'color': RGBColor(0, 0, 0), 'bold': True}
    }
    
    for style_name, properties in heading_styles.items():
        style = doc.styles[style_name]
        style.font.size = Pt(properties['size'])
        style.font.color.rgb = properties['color']
        style.font.bold = properties['bold']
        # Set outline level for SEO
        style.paragraph_format.outline_level = int(style_name.split()[-1])

def add_image(doc, image_url, caption=None):
    """Add an image to the document with optional caption."""
    try:
        # Download image
        response = requests.get(image_url)
        if response.status_code == 200:
            image_stream = BytesIO(response.content)
            
            # Add image
            doc.add_picture(image_stream, width=Inches(6.0))
            
            # Add caption if provided
            if caption:
                caption_para = doc.add_paragraph()
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption_para.add_run(caption)
                caption_run.font.size = Pt(10)
                caption_run.font.italic = True
            
            # Add spacing after image
            doc.add_paragraph()
    except Exception as e:
        print(f"Error adding image {image_url}: {str(e)}")
        # Add error message in document
        doc.add_paragraph(f"[Error: Could not load image from {image_url}]")

async def generate_word_document(sessionId: str):
    """Generates a Word document from session content with proper formatting and images."""
    try:
        session_object_id = ObjectId(sessionId)
    except:
        raise HTTPException(status_code=400, detail="Invalid session ID format")

    # Get the session
    session = await db.sessions.find_one({"_id": session_object_id})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    # Create a new Word document
    doc = Document()
    
    # Set up heading styles
    set_heading_style(doc)
    
    # Add title
    title = doc.add_heading(session.get('title', 'Session Documentation'), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    metadata = doc.add_paragraph()
    metadata.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    metadata.add_run(f"Session ID: {sessionId}\n")
    
    # Get all artifacts for this session in chronological order
    artifacts = await db.artifacts.find(
        {"sessionId": session_object_id}
    ).sort("captureDate", 1).to_list(100)

    if not artifacts:
        raise HTTPException(status_code=404, detail="No artifacts found for this session")

    # Process each artifact
    for artifact in artifacts:
        # Get the latest update for this artifact
        latest_update = await db.artifact_updates.find_one(
            {"artifactId": artifact["_id"]},
            sort=[("processedAt", -1)]
        )
        
        if latest_update and "processedText" in latest_update:
            if artifact["captureType"] == "image":
                # Add image with caption
                caption = f"Image captured at {artifact['captureDate'].strftime('%Y-%m-%d %H:%M:%S')}"
                add_image(doc, artifact["url"], caption)
            else:
                # Add text content with appropriate heading
                heading = doc.add_heading(f"{artifact['captureType'].title()} Content", level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Add timestamp
                timestamp = doc.add_paragraph()
                timestamp.add_run(f"Captured at: {artifact['captureDate'].strftime('%Y-%m-%d %H:%M:%S')}\n")
                
                # Add content
                content = doc.add_paragraph()
                content.add_run(latest_update["processedText"])

    # Create exports directory in the project root
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    exports_dir = os.path.join(base_dir, "exports")
    os.makedirs(exports_dir, exist_ok=True)
    
    # Generate filename
    filename = f"session_{sessionId}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(exports_dir, filename)
    
    # Save the document
    doc.save(filepath)
    
    return {
        "message": "Session exported successfully",
        "filename": filename,
        "filepath": filepath
    } 