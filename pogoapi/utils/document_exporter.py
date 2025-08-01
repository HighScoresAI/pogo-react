"""
Utility for exporting documentation to Word documents.
"""

import os
from typing import Optional, List, Dict, Any
from docx import Document
from docx.shared import Inches
from pogoapi.utils.mongodb_client import MongoDBClient
from pogoapi.utils.path_utils import PathBuilder
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentExporter:
    """
    Handles exporting documentation to Word documents.
    """
    
    def __init__(self, db_client: MongoDBClient, path_builder: PathBuilder):
        """
        Initialize with MongoDB client and PathBuilder.
        
        Args:
            db_client: MongoDB client instance
            path_builder: PathBuilder instance
        """
        self.db_client = db_client
        self.path_builder = path_builder
        
    def export_session_document(self, session_id: str) -> Optional[str]:
        """
        Export a session document with all processed updates
        
        Args:
            session_id: The session ID to export
            
        Returns:
            Optional[str]: Path to the exported document or None if export fails
        """
        try:
            # Get all processed updates for the session
            updates = self.db_client.get_session_updates(session_id)
            
            if not updates:
                print(f"No processed updates found for session {session_id}")
                return None
                
            # Filter for image analysis updates
            image_updates = [u for u in updates if u["update_type"] == "image_analysis"]
            
            if not image_updates:
                print(f"No processed image updates found for session {session_id}")
                return None
                
            # Create document with image updates
            doc = Document()
            doc.add_heading(f"Session Analysis - {session_id}", 0)
            
            for update in image_updates:
                self._add_image_analysis_section(doc, update)
                
            # Get session folder path
            session_path = self.path_builder.build_session_path(session_id)
            
            # Save the document in the session folder
            output_path = os.path.join(session_path, f"session_{session_id}.docx")
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            print(f"Error exporting session document: {e}")
            return None
            
    def _add_audio_section(self, doc: Document, update: Dict[str, Any]) -> None:
        """
        Add an audio section to the document.
        
        Args:
            doc: Word document to add to
            update: Audio update data
        """
        try:
            doc.add_heading('Audio Transcript', level=1)
            doc.add_paragraph(update["content"])
            doc.add_paragraph()  # Add spacing
            
        except Exception as e:
            logger.error(f"Error adding audio section: {str(e)}")
            
    def _add_image_section(self, doc: Document, update: Dict[str, Any]) -> None:
        """
        Add an image section to the document.
        
        Args:
            doc: Word document to add to
            update: Image update data
        """
        try:
            # Get artifact details
            artifact = self.db_client.get_documentation(
                {"_id": update["artifactId"]},
                "artifacts"
            )
            if not artifact:
                logger.error(f"Artifact not found: {update['artifactId']}")
                return
                
            # Add image if available
            image_path = self.path_builder.build_artifact_path(artifact["url"])
            if os.path.exists(image_path):
                doc.add_picture(image_path, width=Inches(6))
                
            # Add analysis
            doc.add_heading('Image Analysis', level=2)
            doc.add_paragraph(update["content"])
            doc.add_paragraph()  # Add spacing
            
        except Exception as e:
            logger.error(f"Error adding image section: {str(e)}")

    def _add_image_analysis_section(self, doc: Document, update: Dict[str, Any]) -> None:
        """
        Add an image analysis section to the document.
        
        Args:
            doc: Word document to add to
            update: Image analysis update data
        """
        try:
            # Get artifact details
            artifact = self.db_client.get_documentation(
                {"_id": update["artifactId"]},
                "artifacts"
            )
            if not artifact:
                logger.error(f"Artifact not found: {update['artifactId']}")
                return
                
            # Add image if available
            image_path = self.path_builder.build_artifact_path(artifact["url"])
            if os.path.exists(image_path):
                doc.add_picture(image_path, width=Inches(6))
                
            # Add analysis
            doc.add_heading('Image Analysis', level=2)
            doc.add_paragraph(update["content"])
            doc.add_paragraph()  # Add spacing
            
        except Exception as e:
            logger.error(f"Error adding image analysis section: {str(e)}") 